from docx import Document
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import argparse

dir_path = os.getcwd()+'\source_code'


def run(args):

    document = Document(args.template)
    dir_path = args.src

    head = document.sections[0].header
    head.paragraphs[0].text = args.output.split('.')[0] + ' v' + args.version

    file_list = []

    for root, dirs, files in os.walk(dir_path):
        for file in files:
            file_list.append(os.path.relpath(
                os.path.join(root, file), dir_path).split('\\'))

    black_list = args.black_list

    for i in file_list:
        with open(os.path.join(dir_path, *i), 'r', encoding='utf-8') as f:
            if i[-1].split('.')[-1] not in black_list:
                # document.add_heading('/'.join(i), 1)

                lines = f.readlines()
                content = '源代码文件路径：'+'/'.join(i)+'\n'
                for line in lines:
                    frist_ch = ''
                    for j in range(len(line)):
                        if line[j] == ' ':
                            continue
                        else:
                            frist_ch = line[j]
                            break
                    if frist_ch == '/':
                        continue
                    if frist_ch == '#':
                        continue
                    if frist_ch == '\n':
                        continue

                    content += line
                doc_p = document.add_paragraph(content)
                doc_p.paragraph_format.line_spacing = 1
                doc_p.style.font.size = docx.shared.Pt(10)

    document.save(args.output)


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--template', type=str,
                        default='template.docx', required=False)
    parser.add_argument('--src', type=str, default=dir_path, required=False)
    parser.add_argument('--output', type=str,
                        default='output.docx', required=False)
    parser.add_argument('--version', type=str, default='1.0', required=False)
    parser.add_argument('--black_list', type=list,
                        default=['png', 'jpg'], required=False)
    args = parser.parse_args()

    run(args)
    print('Done!')
